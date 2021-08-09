import discord
from docx import Document
from docx.shared import Inches
import os
async def fetchannel(message):
  messages = await message.channel.history(limit=None).flatten()
  context = [str(i.content) for i in messages]
  with open("file.txt", "w") as output:
    output.write(convertstring(context))
def convertstring(list):
  
  list.reverse()
  changelist=[str(i).replace("'","").replace("'","").replace("(","").replace(")","").lstrip().replace(",","").replace("!export","").replace("List show all","").replace("successfully add","").replace("List add","").lstrip() for i in list]
  
  listostring = "\n".join(changelist)
  return listostring
async def convertdocx(message):
  messages = await message.channel.history(limit=None).flatten()
  messages.reverse()
  document = Document()
  for i in messages:
    if len(i.content)>0:
      inside = str(i.content).replace("'","").replace("'","").replace("(","").replace(")","").lstrip().replace(",","").replace("!export","").replace("List show all","").replace("successfully add","").replace("List add","").lstrip()
      document.add_paragraph(inside)
    if len(i.attachments) > 0:
        attachment = i.attachments[0]
        if attachment.filename.endswith(".jpg") or attachment.filename.endswith(".jpeg") or attachment.filename.endswith(".png") or attachment.filename.endswith(".webp") or attachment.filename.endswith(".gif"):
            await attachment.save('1.png')
            document.add_picture('1.png')
            os.remove('1.png')

  document.add_page_break()
  document.save('demo.docx')


